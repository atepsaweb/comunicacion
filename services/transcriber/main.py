import asyncio
import os
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from faster_whisper import WhisperModel

MODEL_SIZE = os.getenv("WHISPER_MODEL", "medium")
COMPUTE_TYPE = os.getenv("WHISPER_COMPUTE_TYPE", "int8")
MAX_CONCURRENT = int(os.getenv("MAX_CONCURRENT_JOBS", "2"))

app = FastAPI()
_semaphore: Optional[asyncio.Semaphore] = None
_model: Optional[WhisperModel] = None


def get_model() -> WhisperModel:
    global _model
    if _model is None:
        _model = WhisperModel(MODEL_SIZE, device="cpu", compute_type=COMPUTE_TYPE)
    return _model


@app.on_event("startup")
async def startup() -> None:
    global _semaphore
    _semaphore = asyncio.Semaphore(MAX_CONCURRENT)
    # Warm up: cargar el modelo en memoria antes del primer request
    get_model()


# ─── Audio ────────────────────────────────────────────────────────────────────

class TranscribeRequest(BaseModel):
    path: str


class TranscribeResponse(BaseModel):
    text: str
    duration_sec: int
    model: str


@app.get("/health")
def health() -> dict:
    return {"status": "ok", "model": MODEL_SIZE}


@app.post("/transcribe", response_model=TranscribeResponse)
async def transcribe(req: TranscribeRequest) -> TranscribeResponse:
    path = Path(req.path)
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"File not found: {req.path}")
    if not path.is_file():
        raise HTTPException(status_code=400, detail="Path is not a file")

    assert _semaphore is not None, "Semaphore not initialized"
    async with _semaphore:
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, _do_transcribe, str(path))

    return result


def _do_transcribe(path: str) -> TranscribeResponse:
    model = get_model()
    segments, info = model.transcribe(
        path,
        language="es",
        beam_size=5,
        vad_filter=True,
        vad_parameters={"min_silence_duration_ms": 500},
    )
    text = " ".join(seg.text.strip() for seg in segments).strip()
    duration_sec = int(info.duration)
    return TranscribeResponse(text=text, duration_sec=duration_sec, model=MODEL_SIZE)


# ─── Documentos (PDF y Word) ──────────────────────────────────────────────────

SUPPORTED_DOC_MIMES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}


class ExtractDocumentRequest(BaseModel):
    path: str
    mime_type: str


class ExtractDocumentResponse(BaseModel):
    text: str
    mime_type: str
    chars: int


@app.post("/extract-document", response_model=ExtractDocumentResponse)
async def extract_document(req: ExtractDocumentRequest) -> ExtractDocumentResponse:
    path = Path(req.path)
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"File not found: {req.path}")
    if not path.is_file():
        raise HTTPException(status_code=400, detail="Path is not a file")
    if req.mime_type not in SUPPORTED_DOC_MIMES:
        raise HTTPException(status_code=422, detail=f"Unsupported mime type: {req.mime_type}")

    assert _semaphore is not None, "Semaphore not initialized"
    async with _semaphore:
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, _do_extract_document, str(path), req.mime_type)

    return ExtractDocumentResponse(text=text, mime_type=req.mime_type, chars=len(text))


def _do_extract_pdf(path: str) -> str:
    """Extrae texto de un PDF usando pdfplumber."""
    import pdfplumber  # importación diferida para no romper el startup si no está instalado
    text_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())
    return "\n\n".join(text_parts)


def _do_extract_docx(path: str) -> str:
    """Extrae texto de un archivo Word (.docx)."""
    import docx  # python-docx, importación diferida
    doc = docx.Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # También extraer texto de tablas
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _do_extract_document(path: str, mime_type: str) -> str:
    if mime_type == "application/pdf":
        return _do_extract_pdf(path)
    else:
        # .docx y .doc legacy
        return _do_extract_docx(path)
